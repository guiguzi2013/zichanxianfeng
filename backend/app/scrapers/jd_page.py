"""京东债权页面抓取器（2026-09-01 新增）

背景：queryProductDescription 接口返回的 HTML 只有公告正文（可能极简），但
浏览器渲染后的页面还包含：
  1. 「标的物属性」区块：债权本金/有无抵质押物/担保方式/抵质押物类型/标的物所在地（结构化）
  2. 「附件下载」区：storage.jd.com 上的 .docx/.pdf 附件（资产清单/判决书/承诺书等）

本模块用 Playwright 渲染真实页面，提取属性区块与附件链接；附件按文件名分类，
只下载信息类附件（资产清单/债权明细/抵押物清单/判决书/裁定书等），跳过格式文件
（承诺书/保密函/须知/合同模板），并与债权关联存储。
"""
import logging
import os
import re

logger = logging.getLogger(__name__)

# 重要文件关键词（2026-09-01 用户确认：只保留对尽调有用的重要文件）：
#   - 权属证明类：房产证/土地证明/不动产权
#   - 司法文书类：判决书/裁定书/执行
#   - 资产明细类：xx清单/债权明细/抵押物清单/评估报告
#   - 声明/公告类：声明/处置公告
# 排除：竞买公告/成交确认书/转让协议/合同模板 等交易平台规则类（买后才用）与承诺书类
INFO_ATTACHMENT_KEYWORDS = (
    "房产证", "不动产权", "土地证", "土地证明", "权属",  # 权属证明
    "判决书", "裁定书", "判决", "执行", "裁定",  # 司法文书
    "资产清单", "债权明细", "抵押物清单", "抵质押物清单", "资产明细", "明细表", "债权表", "清单",  # 明细清单
    "评估报告", "评估价", "估价报告", "资产评估", "估值",  # 评估报告
    "声明", "处置公告", "招商公告", "拍卖公告",  # 声明/公告
)
# 交易平台规则/流程类关键词：买后才用到的合同协议注意事项，不下载
SKIP_ATTACHMENT_KEYWORDS = (
    "竞买公告", "成交确认书", "转让协议", "资产转让协议", "合同模板", "债权转让合同",
    "承诺书", "承诺函", "保密", "须知", "确认函", "反洗钱", "涉黑", "委托书", "授权书",
    "登记表", "备忘录", "流程", "指引", "说明", "注意事项",
)


def _classify_attachment(filename: str) -> str:
    """附件分类：info=重要文件(下载解析) / valuation=评估报告 / skip=交易规则/流程类(跳过) / other=其他"""
    name = filename.lower()
    if any(k in filename for k in SKIP_ATTACHMENT_KEYWORDS):
        return "skip"
    if any(k in filename for k in ("评估报告", "评估价", "估价报告", "资产评估", "估值")):
        return "valuation"  # 评估报告：抵押物参考价值（注意年份）
    if any(k in filename for k in INFO_ATTACHMENT_KEYWORDS):
        return "info"
    if name.endswith((".docx", ".doc", ".pdf", ".xlsx", ".xls")):
        return "other"  # 可能是未知文件，暂存不解析（前端不重点展示）
    return "skip"


def _ext_ok(url: str) -> bool:
    """附件 URL 是否是可下载的文件类型（word/excel/pdf 等）"""
    return bool(re.search(r'\.(docx?|pdf|xlsx?|txt)$', url, re.I))


# 标的物属性已知键（2026-09-02：支持表格键值隔行形态，含"抵保方式"等京东变体）
_KNOWN_PROP_KEYS = (
    "债权本金", "有无抵质押物", "担保方式", "抵保方式", "抵质押物类型",
    "抵质押物地址", "标的物所在地", "抵质押物", "保证金", "起拍价",
)


def fetch_jd_page_data(paimai_id: str) -> dict:
    """Playwright 渲染京东债权页，返回 {properties:{}, attachments:[{name,url,type}]}

    properties 字段（页面「标的物详情/标的物属性」区块，键名按页面文本）：
      债权本金/有无抵质押物/担保方式(或"抵保方式")/抵质押物类型/抵质押物地址/标的物所在地 等
    attachments 按 _classify_attachment 分类。

    2026-09-02 修复：属性区可能是**表格结构**（innerText 为"键\\n值"隔行、无冒号，
    如"标的物详情\\n债权本金\\n￥10000000\\n有无抵质押物\\n有\\n抵保方式\\n抵押…"），
    原逻辑只认"键：值"冒号行导致 453 案例标的物详情未抓到；现兼容两种形态。
    """
    from playwright.sync_api import sync_playwright
    from .browser import launch_chromium

    url = f"https://paimai.jd.com/{paimai_id}"
    result: dict = {"properties": {}, "attachments": [], "tables": []}
    try:
        with sync_playwright() as p:
            browser = launch_chromium(p, headless=True)
            page = browser.new_context(viewport={"width": 1440, "height": 2000}).new_page()
            try:
                page.goto(url, wait_until="domcontentloaded", timeout=60000)
                page.wait_for_timeout(8000)  # 等 JS 渲染属性区/附件区
                # 1. 标的物属性区块：页面常为 键:值 列表 或 "标的物详情"表格（键值隔行）
                props: dict = {}

                def _extract_block(block: str) -> None:
                    lines = [ln.strip() for ln in block.split("\n") if ln.strip()]
                    # 形态A：键：值 —— 只接受已知属性键（防整页公告文本被误当属性，2026-09-02 自查修复）
                    for ln in lines:
                        m = re.match(r'^([^：:]+)[：:]\s*(.+)$', ln)
                        if m:
                            key = m.group(1).strip()
                            val = m.group(2).strip()
                            if key in _KNOWN_PROP_KEYS and val:
                                props[key] = val
                    # 形态B：表格键值隔行（键行后紧跟值行；值行不能再是已知键）
                    for i, ln in enumerate(lines):
                        if ln in _KNOWN_PROP_KEYS and i + 1 < len(lines):
                            val = lines[i + 1]
                            if val not in _KNOWN_PROP_KEYS and val not in ("标的物属性", "标的物详情"):
                                props[ln] = val

                try:
                    # 只找含"标的物详情/标的物属性"标题的块（2026-09-02 自查修复：
                    # 原条件"含债权本金+抵质押物"会误中整页公告 div，把公告全文当属性）
                    blocks = page.eval_on_selector_all(
                        "div",
                        """els => els.filter(e => {
                            const t = (e.innerText||'');
                            return (t.includes('标的物详情') || t.includes('标的物属性')) && t.length < 2000;
                        }).map(e => e.innerText)""",
                    )
                    for block in blocks:
                        _extract_block(block)
                except Exception as e:  # noqa: BLE001
                    logger.debug("京东属性区解析失败 %s: %s", paimai_id, e)
                # 形态B 兜底：直接抓页面文本中"标的物详情/标的物属性"后的键值序列
                if not props:
                    try:
                        body_text = page.inner_text("body") if page.locator("body").count() else ""
                        _extract_block(body_text)
                    except Exception:  # noqa: BLE001
                        pass
                result["properties"] = props

                # 3. 页面表格（2026-09-02 修复：此前只抓属性 div 区块，漏掉页面 <table>——
                #    465 案例：债权表格(债务人/债权合计/本金/利息/代垫/担保情况)未抓到）
                try:
                    tables = page.eval_on_selector_all(
                        "table",
                        """els => els.map(e => {
                            const rows = [...e.querySelectorAll('tr')].map(tr =>
                                [...tr.querySelectorAll('th,td')].map(c => (c.innerText||'').trim()));
                            return rows;
                        })""",
                    )
                    parsed_tables: list[dict] = []
                    for t in tables:
                        if not t:
                            continue
                        # 第一行为表头（或含表头关键词的行）
                        header_idx = None
                        for i, r in enumerate(t):
                            joined = " ".join(r)
                            if any(k in joined for k in ("债务人", "债权本金", "债权合计", "本金余额", "担保情况", "标的物")):
                                header_idx = i
                                break
                        if header_idx is None:
                            continue
                        headers = t[header_idx]
                        rows = t[header_idx + 1:]
                        # 过滤空行/出价记录等无关表
                        joined_h = " ".join(headers)
                        if "竞买" in joined_h or "出价" in joined_h or "时间" in joined_h and "价格" in joined_h:
                            continue
                        parsed_tables.append({"headers": headers, "rows": rows})
                    if parsed_tables:
                        result["tables"] = parsed_tables
                except Exception as e:  # noqa: BLE001
                    logger.debug("京东页面表格解析失败 %s: %s", paimai_id, e)

                # 2. 附件链接：storage.jd.com 的 .docx/.pdf 等
                links = page.eval_on_selector_all(
                    "a",
                    """els => els.filter(e => {
                        const h = e.href || '';
                        return h.includes('storage.jd.com') || h.includes('auction.gateway');
                    }).map(e => ({href: e.href, text: (e.innerText||'').trim()}))""",
                )
                seen: set[str] = set()
                for l in links:
                    href = l.get("href") or ""
                    text = l.get("text") or ""
                    if not href or href in seen or not _ext_ok(href):
                        continue
                    seen.add(href)
                    fname = text or href.split("/")[-1]
                    result["attachments"].append({
                        "name": fname,
                        "url": href,
                        "type": _classify_attachment(fname),
                    })
            except Exception as e:  # noqa: BLE001
                logger.warning("京东页面渲染失败 %s: %s", paimai_id, e)
            finally:
                browser.close()
    except Exception as e:  # noqa: BLE001
        logger.exception("Playwright 启动失败: %s", e)
    return result


def download_attachment(url: str, save_dir: str) -> str | None:
    """下载附件到 save_dir，返回本地路径；失败返回 None"""
    import httpx

    os.makedirs(save_dir, exist_ok=True)
    fname = url.split("/")[-1]
    # 文件名清理 + 防重复
    safe = re.sub(r'[^\w.\-]', "_", fname)
    path = os.path.join(save_dir, safe)
    if os.path.exists(path):
        return path
    try:
        r = httpx.get(url, timeout=60, follow_redirects=True,
                      headers={"User-Agent": "Mozilla/5.0"})
        if r.status_code != 200:
            logger.warning("附件下载失败 %s: %s", url, r.status_code)
            return None
        with open(path, "wb") as f:
            f.write(r.content)
        return path
    except Exception as e:  # noqa: BLE001
        logger.exception("附件下载异常 %s: %s", url, e)
        return None


def extract_docx_tables(path: str) -> list[dict]:
    """解析 docx 中的表格 → [{headers:[...], rows:[[...]]}]（python-docx 表格）"""
    try:
        from docx import Document

        doc = Document(path)
        out = []
        for table in doc.tables:
            headers: list[str] = []
            rows: list[list[str]] = []
            for i, row in enumerate(table.rows):
                cells = [c.text.strip() for c in row.cells]
                if i == 0:
                    headers = cells
                else:
                    rows.append(cells)
            if headers or rows:
                out.append({"headers": headers, "rows": rows})
        return out
    except Exception as e:  # noqa: BLE001
        logger.warning("docx 表格解析失败 %s: %s", path, e)
        return []


def extract_excel_tables(path: str) -> list[dict]:
    """解析 excel(xlsx/xls) 表格 → [{headers:[...], rows:[[...]]}]（openpyxl）

    原文附件常为 word/excel/pdf 混合格式（2026-09-01 用户提示）；excel 的资产清单/评估表
    也需解析。xls 旧格式 openpyxl 不支持时降级 pandas。
    """
    try:
        if path.lower().endswith((".xlsx",)):
            from openpyxl import load_workbook

            wb = load_workbook(path, read_only=True, data_only=True)
        else:
            # .xls 旧格式：openpyxl 不支持，尝试 pandas
            import pandas as pd

            sheets = pd.read_excel(path, sheet_name=None, header=None)
            out = []
            for df in sheets.values():
                rows = [[("" if c is None else str(c)).strip() for c in row] for row in df.values.tolist()]
                if not rows:
                    continue
                headers, body = rows[0], rows[1:]
                if headers or body:
                    out.append({"headers": headers, "rows": body})
            return out

        out = []
        for ws in wb.worksheets:
            rows = [[("" if c is None else str(c)).strip() for c in row] for row in ws.iter_rows(values_only=True)]
            rows = [r for r in rows if any(r)]
            if not rows:
                continue
            headers, body = rows[0], rows[1:]
            if headers or body:
                out.append({"headers": headers, "rows": body})
        return out
    except Exception as e:  # noqa: BLE001
        logger.warning("excel 解析失败 %s: %s", path, e)
        return []


def extract_text_from_any(path: str) -> str | None:
    """按扩展名提取文本（docx/pdf/excel/txt），供评估报告内容分析"""
    import os

    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in (".docx",):
            from docx import Document
            doc = Document(path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for row in t.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            return "\n".join(parts)[:20000]
        if ext in (".pdf",):
            from pypdf import PdfReader
            reader = PdfReader(path)
            return "\n".join((page.extract_text() or "") for page in reader.pages)[:20000]
        if ext in (".xlsx", ".xls"):
            tables = extract_excel_tables(path)
            parts = []
            for t in tables:
                parts.append(" | ".join(t.get("headers") or []))
                for r in t.get("rows") or []:
                    parts.append(" | ".join(r))
            return "\n".join(parts)[:20000]
        if ext in (".txt", ".md"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()[:20000]
    except Exception as e:  # noqa: BLE001
        logger.warning("文本提取失败 %s: %s", path, e)
    return None
