"""补充材料解析服务（P1）

上传文件 → 提取文本 → LLM 识别内容类型 → 分发到对应报告版块 → 触发重新生成。

P0/P1 实现分层：
1. 文本提取：txt 直接读；docx/pdf 需要 python-docx/pypdf（requirements 已加）；失败降级标注
2. 类型识别：优先扩展名；其次 LLM（mock 可用）
3. 分发：将识别结果写入 report.supplements（JSON），由 report_builder 在重新生成时引用
"""
import logging
import os
import re

logger = logging.getLogger(__name__)

# RapidOCR 引擎单例（2026-09-05：多文件识别时每份都重新加载模型很慢，只初始化一次复用）
_OCR_ENGINE = None
_OCR_LOCK = None


def _get_ocr():
    """懒加载共享 RapidOCR 引擎（线程安全，首次约数秒加载模型，后续零成本复用）"""
    global _OCR_ENGINE, _OCR_LOCK
    if _OCR_ENGINE is None:
        try:
            from rapidocr_onnxruntime import RapidOCR  # 延迟导入
        except ImportError:
            logger.warning("rapidocr 未安装，图片/扫描件识别不可用")
            return None
        import threading
        _OCR_LOCK = threading.Lock()
        with _OCR_LOCK:
            if _OCR_ENGINE is None:
                _OCR_ENGINE = RapidOCR()
    return _OCR_ENGINE

# 文件类型 → 建议分发版块
TYPE_TARGET = {
    "判决书": ["claim_basic", "legal"],
    "裁定书": ["legal"],
    "执行裁定": ["legal"],
    "评估报告": ["collateral"],
    "评估": ["collateral"],
    "尽调说明": ["risk", "disposal"],
    "情况说明": ["risk"],
    "合同": ["claim_basic"],
    "债权转让": ["claim_basic"],
}


def extract_text_from_file(path: str, file_type: str | None, progress=None) -> str | None:
    """提取文件文本。失败返回 None（调用方标注'无法解析'）。
    docx 段落+表格全收（2026-09-04：Word 里的债权清单表此前丢失，须一并提取）。
    progress: 可选回调 progress(pct: 0-100, phase: str)，供任务按文件显示进度条（2026-09-05）。"""
    try:
        if file_type in ("txt", "text", "md"):
            if progress:
                progress(60, "读取文本")
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()[:20000]
            if progress:
                progress(100, "完成")
            return text
        if file_type in ("docx",):
            return _extract_docx(path, progress)
        if file_type in ("pdf",):
            return _extract_pdf(path, progress)
        if file_type in ("jpg", "jpeg", "png", "webp", "bmp"):
            return _ocr_image(path, progress)
        if file_type in ("xlsx", "xls", "csv"):
            # Excel：逐单元格文本化（表格行以制表符连接）——内容可能是债权列表/抵押物清单等，
            # 交给后续 LLM 按内容识别，不再默认当作债权列表（2026-09-05 用户确认）
            return _extract_excel(path, file_type, progress)
    except Exception as e:  # noqa: BLE001
        logger.warning("extract text failed for %s: %s", path, e)
        return None
    return None


def _extract_excel(path: str, file_type: str, progress=None) -> str | None:
    """Excel/Csv 文本化：所有 sheet 的单元格按行以制表符连接（保留表格结构感）。
    表头与单位随文本保留，供 LLM 识别是债权列表 / 抵押物清单 / 无关内容。"""
    import csv as _csv

    try:
        rows_out: list[str] = []
        if file_type == "csv":
            if progress:
                progress(40, "读取 CSV")
            with open(path, "r", encoding="utf-8", errors="ignore") as fp:
                for row in _csv.reader(fp):
                    cells = [str(c).strip() for c in row if str(c).strip()]
                    if cells:
                        rows_out.append("\t".join(cells))
        elif file_type == "xlsx":
            from openpyxl import load_workbook

            if progress:
                progress(20, "读取 Excel")
            wb = load_workbook(path, read_only=True, data_only=True)
            for ws in wb.worksheets:
                if len(rows_out) > 500:
                    break
                sheet_name = ws.title or ""
                for raw in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in raw if c is not None and str(c).strip()]
                    if cells:
                        rows_out.append(f"[{sheet_name}] " + "\t".join(cells))
        elif file_type == "xls":
            import xlrd

            if progress:
                progress(20, "读取 Excel(.xls)")
            wb = xlrd.open_workbook(path)
            for ws in wb.sheets():
                if len(rows_out) > 500:
                    break
                for r in range(ws.nrows):
                    cells = [str(ws.cell_value(r, c)).strip() for c in range(ws.ncols) if str(ws.cell_value(r, c)).strip()]
                    if cells:
                        rows_out.append("\t".join(cells))
        if progress:
            progress(100, "完成")
        return "\n".join(rows_out)[:20000] if rows_out else None
    except Exception as e:  # noqa: BLE001
        logger.warning("excel extract failed for %s: %s", path, e)
        return None


def _extract_docx(path: str, progress=None) -> str | None:
    """docx 全文提取：正文段落 + 表格（表格每行以制表符连接，保留表头与结构感）。
    Word 中的债权清单/资产表在此完整捞回，供后续 LLM 按内容识别单户或批量。"""
    try:
        from docx import Document  # 延迟导入

        doc = Document(path)
        parts: list[str] = []
        if progress:
            progress(30, "读取正文")
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        if progress:
            progress(70, "读取表格")
        for tb in doc.tables:
            for row in tb.rows:
                cells = []
                for c in row.cells:
                    txt = (c.text or "").strip()
                    if txt and txt not in cells:  # 合并单元格会重复同一文本，去重防错位
                        cells.append(txt)
                if cells:
                    parts.append("\t".join(cells))
        if progress:
            progress(100, "完成")
        return "\n".join(parts)[:20000] if parts else None
    except Exception as e:  # noqa: BLE001
        logger.warning("docx extract failed for %s: %s", path, e)
        return None


def _extract_pdf(path: str, progress=None) -> str | None:
    """PDF 文本提取：优先文字层；扫描版（文字不足）时逐页渲染 OCR，支持多页。"""
    from pypdf import PdfReader  # 延迟导入

    if progress:
        progress(15, "读取 PDF 文字层")
    try:
        reader = PdfReader(path)
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as e:  # noqa: BLE001
        logger.warning("pypdf failed for %s: %s", path, e)
        text = ""
    if text and len(text.strip()) >= 100:
        if progress:
            progress(100, "完成")
        return text[:20000]
    # 扫描版 PDF：逐页转图 OCR
    if progress:
        progress(20, "扫描件，逐页 OCR 中…")
    ocr_text = _ocr_pdf_pages(path, progress)
    if progress:
        progress(100, "完成")
    return ocr_text[:20000] if ocr_text else (text[:20000] if text else None)


def _ocr_pdf_pages(path: str, progress=None) -> str | None:
    """扫描版 PDF 逐页渲染 + OCR（PyMuPDF + RapidOCR，引擎复用单例）。
    progress(pct) 按已处理页数/总页数更新（50 页以内区间，页数多时收敛到 95 避免误导）。"""
    import os
    import uuid

    try:
        import pymupdf  # PyMuPDF
    except ImportError:
        logger.warning("pymupdf 未安装，扫描版 PDF 无法识别")
        return None
    ocr = _get_ocr()
    if ocr is None:
        return None
    try:
        doc = pymupdf.open(path)
        parts: list[str] = []
        total_pages = len(doc) or 1
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=150)  # 150dpi 兼顾速度与识别率
            img_path = os.path.join(os.path.dirname(path), f"_ocr_{uuid.uuid4().hex}.png")
            try:
                pix.save(img_path)
                with _OCR_LOCK:  # onnxruntime 引擎并发调用不安全，串行保护
                    result, _ = ocr(str(img_path))
                if result:
                    parts.append("\n".join(text for _, text, _ in result))
            finally:
                if os.path.exists(img_path):
                    try:
                        os.remove(img_path)
                    except OSError:
                        pass
            if i >= 30:  # 页数上限保护
                break
            if progress:
                # 20% → 95%：随已处理页数推进（上限30页内线性，超30页压缩区间）
                span = min(total_pages, 30)
                progress(round(20 + 75 * (i + 1) / max(span, 1)), f"扫描件 OCR 第 {i + 1}/{min(total_pages, 30)} 页")
        return "\n".join(parts)[:20000] if parts else None
    except Exception as e:  # noqa: BLE001
        logger.warning("PDF OCR failed for %s: %s", path, e)
        return None


def _ocr_image(path: str, progress=None) -> str | None:
    """图片 OCR（RapidOCR 离线引擎，中文支持；引擎复用单例）。未安装时返回 None。"""
    ocr = _get_ocr()
    if ocr is None:
        return None
    try:
        if progress:
            progress(30, "图片识别中…")
        with _OCR_LOCK:
            result, _ = ocr(str(path))
        if progress:
            progress(100, "完成")
        if result:
            return "\n".join(text for _, text, _ in result)[:20000]
    except Exception as e:  # noqa: BLE001
        logger.warning("OCR failed for %s: %s", path, e)
    return None


def classify_content(text: str | None, filename: str = "") -> dict:
    """识别文件内容类型与目标版块。

    Returns:
        {"file_type": "判决书|评估报告|...", "target_sections": [...], "summary": "..."}
    """
    # 文件名关键词优先
    name_hint = ""
    for kw, target in TYPE_TARGET.items():
        if kw in filename:
            name_hint = kw
            break

    if text:
        for kw, target in TYPE_TARGET.items():
            if kw in text[:500]:
                return {
                    "file_type": kw,
                    "target_sections": target,
                    "summary": _make_summary(text),
                }

    if name_hint:
        return {
            "file_type": name_hint,
            "target_sections": TYPE_TARGET[name_hint],
            "summary": f"（文件名识别）{filename}",
        }

    return {
        "file_type": "其他材料",
        "target_sections": ["risk"],
        "summary": _make_summary(text) if text else "无法解析内容，请人工查看",
    }


def _make_summary(text: str, max_len: int = 300) -> str:
    """生成内容摘要（取前 N 字，去空白）"""
    clean = re.sub(r"\s+", " ", text or "").strip()
    return clean[:max_len] + ("…" if len(clean) > max_len else "")
