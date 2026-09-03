"""土地价格库解析服务：从文本/Excel/Word 自动提取 地区/土地性质/单价 并归类

支持：
- 粘贴文本：逐行/逐段识别『山东-青岛 工业 600~1200元/㎡』『青岛 工业用地 30万/亩』等
- Excel 导入：表头同义词识别（地区/城市/土地性质/单价/上限/下限等）
- Word 导入：抽取文本后走文本解析
"""
import json
import logging
import re

logger = logging.getLogger(__name__)

# 土地性质同义词 → 统一归类（对齐《土地利用现状分类》GB/T 21010-2017 一级类）
LAND_TYPE_SYNONYMS = {
    "工业": ["工业", "工矿", "工业用地", "工矿仓储", "厂房用地", "工业仓储", "产业用地", "工业房地产"],
    "商业": ["商业", "商服", "商业服务业", "商住", "商业用地", "商服用地", "批发零售", "商务金融"],
    "住宅": ["住宅", "居住", "住宅用地", "居住用地", "商品房", "城镇住宅"],
    "综合": ["综合", "其他", "混合", "综合用地"],
    "仓储": ["仓储", "物流", "仓储用地", "物流仓储"],
    "农业": ["农业", "耕地", "园地", "林地", "草地", "设施农用地", "农用地"],
    "公共": ["公共", "教育", "医疗", "机关团体", "公共服务"],
    "交通": ["交通", "运输", "交通运输", "铁路", "公路", "机场", "港口"],
}

# 中国省份列表（用于识别）
PROVINCES = [
    "北京", "天津", "上海", "重庆", "河北", "山西", "辽宁", "吉林", "黑龙江",
    "江苏", "浙江", "安徽", "福建", "江西", "山东", "河南", "湖北", "湖南",
    "广东", "海南", "四川", "贵州", "云南", "陕西", "甘肃", "青海", "台湾",
    "内蒙古", "广西", "西藏", "宁夏", "新疆", "香港", "澳门",
]

# Excel 表头同义词
HEADER_SYNONYMS = {
    "province": ["省", "省份", "省/市", "所属省"],
    "city": ["城市", "市", "地区", "区域", "所在城市", "地市", "城市/区"],
    "district": ["区县", "区", "县", "县级", "区/县"],
    "land_type": ["土地性质", "土地类型", "用途", "用地性质", "土地用途", "类型", "用途类型", "性质", "规划用途"],
    "price_lo": ["单价下限", "最低价", "下限", "低价", "最低单价", "保守价", "低值"],
    "price_hi": ["单价上限", "最高价", "上限", "高价", "最高单价", "乐观价", "高值"],
    "price": ["单价", "价格", "均价", "元/㎡", "元/平方米", "每亩价格", "出让价", "基准价"],
    "source": ["来源", "数据来源", "出处"],
    "effective_date": ["生效日期", "日期", "时间", "年份", "年度"],
    "note": ["备注", "说明", "注释"],
}

# Excel 表头同义词（宽松：包含即匹配）
def _match_header(name: str) -> str | None:
    n = (name or "").strip().lower()
    for field, syns in HEADER_SYNONYMS.items():
        for s in syns:
            if s.lower() in n or n in s.lower():
                return field
    return None


def _normalize_land_type(text: str) -> str | None:
    """识别土地性质并归一化"""
    t = text or ""
    for norm, syns in LAND_TYPE_SYNONYMS.items():
        for s in syns:
            if s in t:
                return norm
    return None


def _normalize_region(text: str) -> tuple[str | None, str | None, str | None]:
    """识别 省/市/区县。返回 (province, city, district)"""
    t = text or ""
    province = None
    for p in PROVINCES:
        if t.startswith(p):
            province = p
            rest = t[len(p):].lstrip("-—/· ")
            break
    else:
        rest = t
    # 市：如 山东-青岛 / 青岛市 / 青岛 城阳
    city = None
    district = None
    m = re.search(r"([\u4e00-\u9fff]{2,10}?市(?!场|政|民|中|心))", rest)
    if m:
        city = m.group(1).rstrip("市")
    if not city:
        # 剩余部分开头可能是 区县 或 市名（如 山东-青岛 工业用地 → 青岛；临沂市兰山区 → 临沂）
        m_dist = re.match(r"([\u4e00-\u9fff]{2,8}?)([\u4e00-\u9fff]{2,8}?(?:区|县))", rest)
        if m_dist:
            city = m_dist.group(1)
            district = m_dist.group(2)
        else:
            m_head = re.match(r"([\u4e00-\u9fff]{2,8}?)(?:[\s，,、区县]|$)", rest)
            if m_head and not re.search(r"(?:区|县)$", m_head.group(1)) and m_head.group(1) not in ("城区", "工业", "商业", "住宅", "综合"):
                city = m_head.group(1)
    # 区县（补充：描述中直接带区县）
    if not district:
        m3 = re.search(r"([\u4e00-\u9fff]{2,8}?(?:区|县))(?!市)", t)
        if m3:
            district = m3.group(1)
    return province, city, district


def _parse_price(text: str) -> tuple[int | None, int | None] | None:
    """解析单价区间，返回 (lo, hi) 元/㎡；支持元/㎡、万元/亩"""
    t = text or ""
    # 万元/亩 → 元/㎡（1亩=666.67㎡，1万元=10000元）；支持 30~50万/亩 与 30万/亩~50万/亩
    if "亩" in t:
        # 区间：30~50万/亩
        m_pair = re.search(r"(\d+(?:\.\d+)?)\s*[~—-]\s*(\d+(?:\.\d+)?)\s*万?元?/亩", t)
        if m_pair:
            lo = int(float(m_pair.group(1)) * 10000 / 666.67)
            hi = int(float(m_pair.group(2)) * 10000 / 666.67)
            return min(lo, hi), max(lo, hi)
        # 区间：30万/亩~50万/亩
        nums = re.findall(r"(\d+(?:\.\d+)?)\s*万?元?/亩", t)
        if nums:
            vals = sorted(float(x) for x in nums[:2])
            lo = int(vals[0] * 10000 / 666.67)
            hi = int(vals[-1] * 10000 / 666.67)
            return lo, hi
        nums2 = re.findall(r"(\d+(?:\.\d+)?)\s*万?元?", t)
        if nums2:
            vals = sorted(float(x) for x in nums2[:2])
            lo = int(vals[0] * 10000 / 666.67)
            hi = int(vals[-1] * 10000 / 666.67)
            return lo, hi
    # 元/㎡ 区间
    m = re.search(r"(\d+(?:\.\d+)?)\s*[~—-]\s*(\d+(?:\.\d+)?)\s*(?:元/㎡|元/平方米|元每平|元/平)", t)
    if m:
        return int(float(m.group(1))), int(float(m.group(2)))
    m2 = re.search(r"(\d+(?:\.\d+)?)\s*(?:元/㎡|元/平方米|元每平|元/平)", t)
    if m2:
        v = int(float(m2.group(1)))
        return v, v
    # 裸数字（可能是单价）
    m3 = re.search(r"单价[：:]\s*(\d+(?:\.\d+)?)\s*[~—-]\s*(\d+(?:\.\d+)?)", t)
    if m3:
        return int(float(m3.group(1))), int(float(m3.group(2)))
    return None


def parse_text_lines(text: str) -> list[dict]:
    """解析粘贴文本/Word 文本，返回记录列表 [{province,city,district,land_type,price_lo,price_hi,source,note}]"""
    records = []
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    # 每行/每段独立解析
    for line in lines:
        rec = parse_text_entry(line)
        if rec:
            records.append(rec)
    return records


def parse_text_entry(text: str) -> dict | None:
    """解析单条文本记录"""
    t = text.strip()
    if not t or len(t) < 4:
        return None
    province, city, district = _normalize_region(t)
    land_type = _normalize_land_type(t)
    price = _parse_price(t)
    if not land_type or not price:
        return None
    lo, hi = price
    return {
        "province": province,
        "city": city,
        "district": district,
        "land_type": land_type,
        "price_lo": lo,
        "price_hi": hi,
        "source": "人工录入/文档导入",
        "note": t[:200],
    }


def parse_excel_rows(rows: list[dict]) -> list[dict]:
    """解析 Excel 行（rows 为 {表头: 值} 字典列表），返回记录列表"""
    records = []
    for row in rows:
        mapped = {}
        for header, value in row.items():
            field = _match_header(header)
            if field and value not in (None, ""):
                mapped[field] = str(value)
        rec = _build_record(mapped)
        if rec:
            records.append(rec)
    return records


def _build_record(mapped: dict) -> dict | None:
    """从映射字段构建记录"""
    land_type = _normalize_land_type(mapped.get("land_type", ""))
    if not land_type:
        return None
    province, city, district = _normalize_region(f"{mapped.get('province', '')}{mapped.get('city', '')}{mapped.get('district', '')}")
    lo, hi = None, None
    if mapped.get("price_lo"):
        try:
            lo = int(float(mapped["price_lo"]))
        except ValueError:
            lo = None
    if mapped.get("price_hi"):
        try:
            hi = int(float(mapped["price_hi"]))
        except ValueError:
            hi = None
    if lo is None and hi is None and mapped.get("price"):
        parsed = _parse_price(mapped["price"])
        if parsed:
            lo, hi = parsed
    if lo is None and hi is None:
        return None
    if lo is None:
        lo = hi
    if hi is None:
        hi = lo
    return {
        "province": province or mapped.get("province"),
        "city": city or mapped.get("city"),
        "district": district or mapped.get("district"),
        "land_type": land_type,
        "price_lo": lo,
        "price_hi": hi,
        "source": mapped.get("source") or "人工录入/文档导入",
        "effective_date": mapped.get("effective_date"),
        "note": mapped.get("note"),
    }


def extract_text_from_docx(path: str) -> str:
    """从 Word 提取文本（复用 python-docx）"""
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)
